from docx import Document
import os

docs_to_check = [
    f for f in os.listdir('.')
    if f.startswith('Auditoría_Web') and f.endswith('.docx')
]

print(f"Documentos encontrados: {docs_to_check}\n")

for doc_file in sorted(docs_to_check)[-2:]:
    print(f"\n{'='*60}\nDOCUMENTO: {doc_file}\n{'='*60}")
    doc = Document(doc_file)
    
    # Buscar sección de 4.1 (Sesión autenticada)
    in_auth_section = False
    auth_content = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if '4.1' in text or 'sesión autenticada' in text.lower():
            in_auth_section = True
            print(f"\n[SECCIÓN 4.1 - SESIÓN AUTENTICADA]")
        
        if in_auth_section:
            if text and len(text) > 3:
                print(f"  {text[:140]}")
            if i > 150 and '4.2' in text:
                break
            if i > 180:
                break
    
    # Buscar referencias a cookie
    print(f"\n[BÚSQUEDA DE COOKIE/SESIÓN EN DOCUMENTO]")
    cookie_count = 0
    for i, para in enumerate(doc.paragraphs):
        if 'cookie' in para.text.lower() or 'session' in para.text.lower() or 'sesión' in para.text.lower():
            print(f"  L{i}: {para.text.strip()[:140]}")
            cookie_count += 1
    if cookie_count == 0:
        print("  [SIN REFERENCIAS A COOKIE/SESIÓN]")

print("\n\nANÁLISIS DE HALLAZGOS:")
for doc_file in sorted(docs_to_check)[-2:]:
    print(f"\n--- {doc_file} ---")
    doc = Document(doc_file)
    
    # Contar hallazgos por tipo
    findings = {}
    for para in doc.paragraphs:
        if 'Sev.' in para.text or 'Severidad' in para.text:
            continue
        
        for keyword in ['XSS', 'CSRF', 'SQL', 'RCE', 'Path', 'Open', 'Admin', 'Weak']:
            if keyword.lower() in para.text.lower():
                findings[keyword] = findings.get(keyword, 0) + 1
    
    if findings:
        for key, count in sorted(findings.items(), key=lambda x: -x[1])[:5]:
            print(f"  {key}: {count} referencias")
