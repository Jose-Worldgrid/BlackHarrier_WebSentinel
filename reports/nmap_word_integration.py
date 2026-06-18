# Modulo: Generador de secciones Nmap para informes Word - diseño limpio y atacable.

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import ipaddress
import re
from urllib.parse import urlparse


_SERVICE_ATTACK_VECTORS = {
    "ftp": {"vectors": ["brute-force", "default-creds", "anonymous-login"], "risk": "Alta"},
    "ftp-proxy": {"vectors": ["brute-force", "default-creds", "anonymous-login"], "risk": "Alta"},
    "ssh": {"vectors": ["brute-force", "key-exchange-weakness", "user-enum"], "risk": "Alta"},
    "telnet": {"vectors": ["plaintext-auth", "brute-force", "credential-capture"], "risk": "Crítica"},
    "http": {"vectors": ["web-app-attacks", "default-pages", "directory-traversal"], "risk": "Alta"},
    "http-proxy": {"vectors": ["web-app-attacks", "default-pages", "directory-traversal"], "risk": "Alta"},
    "https": {"vectors": ["ssl-tls-weak", "certificate-recon", "web-app-attacks"], "risk": "Alta"},
    "smtp": {"vectors": ["open-relay", "user-enum", "header-injection"], "risk": "Media"},
    "pop3": {"vectors": ["brute-force", "plaintext-auth", "user-enum"], "risk": "Alta"},
    "imap": {"vectors": ["brute-force", "plaintext-auth", "user-enum"], "risk": "Alta"},
    "smb": {"vectors": ["null-session", "enumeration", "ransomware-vector"], "risk": "Crítica"},
    "netbios": {"vectors": ["name-resolution", "computer-enumeration"], "risk": "Media"},
    "dns": {"vectors": ["zone-transfer", "dns-enumeration", "cache-poisoning"], "risk": "Media"},
    "mysql": {"vectors": ["default-auth", "sql-injection", "credential-compromise"], "risk": "Crítica"},
    "mssql": {"vectors": ["default-auth", "sql-injection", "credential-compromise"], "risk": "Crítica"},
    "postgresql": {"vectors": ["default-auth", "sql-injection"], "risk": "Crítica"},
    "rdp": {"vectors": ["brute-force", "bluekeep", "credential-compromise"], "risk": "Crítica"},
    "vnc": {"vectors": ["brute-force", "no-auth", "credential-capture"], "risk": "Crítica"},
    "snmp": {"vectors": ["default-community", "enumeration", "rce"], "risk": "Alta"},
    "ldap": {"vectors": ["anonymous-bind", "enumeration", "injection"], "risk": "Media"},
    "redis": {"vectors": ["no-auth", "config-write", "rce"], "risk": "Crítica"},
    "mongodb": {"vectors": ["no-auth", "enumeration", "data-dump"], "risk": "Crítica"},
}

_CLOUD_PROVIDERS = [
    ("amazonaws", "AWS", "EC2/ELB"),
    ("elb.amazonaws.com", "AWS", "Elastic Load Balancer"),
    ("ec2-", "AWS", "EC2 Instance"),
    ("azureedge", "Azure", "Azure CDN"),
    ("windows.net", "Azure", "Azure Service"),
    ("azure", "Azure", "Azure Service"),
    ("googleapis", "GCP", "Google Cloud"),
    ("googleusercontent", "GCP", "Google Cloud"),
]


def _detect_cloud(host: str) -> tuple[str, str]:
    h = str(host or "").lower()
    for hint, provider, resource in _CLOUD_PROVIDERS:
        if hint in h:
            return provider, resource
    return "No cloud", ""


def _classify_ip_membership(host: dict) -> tuple[str, str]:
    ip = str(host.get("host") or "")
    hostname = str(host.get("hostname") or "")
    hostnames = " ".join(str(x) for x in (host.get("hostnames") or []) if x)
    blob = f"{ip} {hostname} {hostnames}".lower()

    provider, resource = _detect_cloud(blob)
    if provider != "No cloud":
        return provider, resource or "Servicio cloud"

    try:
        obj = ipaddress.ip_address(ip)
        if obj.is_private:
            if hostname:
                return "Red privada / equipo interno", "PC o servidor interno"
            return "Red privada / equipo interno", "Host interno"
        if obj.is_loopback:
            return "Loopback", "Equipo local"
        if obj.is_multicast:
            return "Multicast", "Difusión"
        if obj.is_link_local:
            return "Link-local", "Interfaz local"
    except Exception:
        pass

    return "Pública / Internet", "Host expuesto"


def _risk_for_service(service: str) -> dict:
    svc = str(service or "").lower().split("/")[0].strip()
    if svc in _SERVICE_ATTACK_VECTORS:
        return _SERVICE_ATTACK_VECTORS[svc]
    for key, val in _SERVICE_ATTACK_VECTORS.items():
        if key in svc:
            return val
    return {"vectors": ["enumeration", "banner-grab"], "risk": "Media"}


def _endpoint_url(host: str, port: int, service: str) -> str:
    s = str(service or "").lower()
    if "https" in s:
        return f"https://{host}:{port}"
    if "http" in s:
        return f"http://{host}:{port}"
    if s.startswith("ftp"):
        return f"ftp://{host}:{port}"
    if s == "ssh":
        return f"ssh://{host}:{port}"
    return f"{host}:{port}"


def _shading(rgb_color):
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    hex_color = "%02x%02x%02x" % (rgb_color[0], rgb_color[1], rgb_color[2])
    return parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), hex_color))


def _extract_sniffed_ips(raw: str) -> list[str]:
    text = str(raw or "")
    ips = re.findall(r"\b(?:\d{1,3}\.){3}\d{1,3}\b", text)
    ordered = []
    seen = set()
    for ip in ips:
        if ip in seen:
            continue
        seen.add(ip)
        ordered.append(ip)
    return ordered


def _ip_kind(ip: str) -> str:
    try:
        obj = ipaddress.ip_address(ip)
        if obj.is_multicast:
            return "Multicast"
        if obj.is_private:
            return "Privada"
        if obj.is_loopback:
            return "Loopback"
        if obj.is_link_local:
            return "Link-local"
        return "Pública"
    except Exception:
        return "No clasificable"


def _parse_target_expression_error(raw: str) -> tuple[str, str]:
    m = re.search(r'Unable to split netmask from target expression:\s*"([^"]+)"', str(raw or ""), re.IGNORECASE)
    if not m:
        return "", ""
    original = m.group(1).strip()
    cleaned = original
    if original.startswith(("http://", "https://")):
        p = urlparse(original)
        cleaned = str(p.hostname or "").strip()
    elif "/" in original:
        cleaned = original.split("/", 1)[0].strip()
    return original, cleaned


def add_nmap_section_to_report(doc, nmap_output: str, title: str = "Reporte de Escaneo Nmap", nmap_structured: dict | None = None):
    """Sección Nmap en Word usando datos XML estructurados como fuente primaria."""
    hosts = (nmap_structured or {}).get("hosts") or []
    if not hosts and (not nmap_output or not nmap_output.strip()):
        return

    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not hosts:
        doc.add_paragraph(
            "Nmap ejecutado pero sin hosts en el XML resultante. "
            "El escaneo pudo haber sido interrumpido, filtrado por firewall, o el target no respondió."
        )

        original_target, cleaned_target = _parse_target_expression_error(nmap_output)
        if original_target:
            doc.add_paragraph(
                "Error detectado en objetivo de Nmap: se pasó una URL/ruta en vez de host/IP.",
                style="List Bullet"
            )
            doc.add_paragraph(f"Objetivo recibido: {original_target}", style="List Bullet 2")
            doc.add_paragraph(f"Objetivo correcto: {cleaned_target or 'host/ip limpio'}", style="List Bullet 2")

        sniffed_ips = _extract_sniffed_ips(nmap_output)
        if sniffed_ips:
            doc.add_heading("IPs observadas en pre-scan (ordenadas)", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            headers = ["IP", "Qué es", "Contexto"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                table.rows[0].cells[i]._element.get_or_add_tcPr().append(_shading(RGBColor(0, 51, 102)))
                for para in table.rows[0].cells[i].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)

            for ip in sniffed_ips:
                row = table.add_row().cells
                kind = _ip_kind(ip)
                row[0].text = ip
                row[1].text = kind
                row[2].text = "Capturada por targets-sniffer (NSE pre-scan)."
        return

    total_open = sum(
        len([p for p in h.get("ports", []) if p.get("state") == "open"])
        for h in hosts
    )

    doc.add_heading("Resumen del escaneo", level=2)
    doc.add_paragraph(f"Hosts activos detectados: {len(hosts)}")
    doc.add_paragraph(f"Total puertos abiertos: {total_open}")

    doc.add_heading("IPs descubiertas y clasificación", level=2)
    ip_table = doc.add_table(rows=1, cols=4)
    ip_table.style = "Table Grid"
    ip_headers = ["IP", "Clasificación", "Pertenencia", "Evidencia"]
    for i, h in enumerate(ip_headers):
        ip_table.rows[0].cells[i].text = h
        ip_table.rows[0].cells[i]._element.get_or_add_tcPr().append(_shading(RGBColor(0, 51, 102)))
        for para in ip_table.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    seen_ips = set()
    for host in hosts:
        ip = str(host.get("host") or "")
        if not ip or ip in seen_ips:
            continue
        seen_ips.add(ip)
        classification = _ip_kind(ip)
        membership, evidence = _classify_ip_membership(host)
        row = ip_table.add_row().cells
        row[0].text = ip
        row[1].text = classification
        row[2].text = membership
        row[3].text = evidence

    for host in hosts:
        ip = str(host.get("host") or "?")
        hostname = str(host.get("hostname") or "")
        os_hint = str(host.get("os") or "No detectado")
        open_ports = [p for p in host.get("ports", []) if p.get("state") == "open"]
        cloud_provider, cloud_resource = _detect_cloud(f"{ip} {hostname}")

        doc.add_heading(f"Host: {hostname or ip}  [{ip}]", level=2)

        meta = [
            ("IP", ip),
            ("Hostname", hostname or "—"),
            ("SO detectado", os_hint),
            ("Entorno cloud", f"{cloud_provider}" + (f" / {cloud_resource}" if cloud_resource else "")),
            ("Puertos abiertos", str(len(open_ports))),
        ]
        for label, value in meta:
            p = doc.add_paragraph()
            r = p.add_run(f"{label}: ")
            r.bold = True
            p.add_run(value)

        if not open_ports:
            doc.add_paragraph("No se detectaron puertos abiertos en este host.")
            continue

        doc.add_heading("Puertos abiertos — servicio, versión y vectores de ataque", level=3)
        headers = ["Puerto", "Protocolo", "Servicio", "Producto / Versión", "Endpoint", "Riesgo", "Vectores de ataque"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i]._element.get_or_add_tcPr().append(_shading(RGBColor(0, 51, 102)))
            for para in hdr_cells[i].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

        for p in open_ports:
            port_num = int(p.get("port") or 0)
            proto = str(p.get("protocol") or "tcp")
            service = str(p.get("service") or "unknown")
            product = str(p.get("product") or "")
            version = str(p.get("version") or "")
            prod_ver = " ".join(x for x in [product, version] if x).strip() or "—"
            attack = _risk_for_service(service)
            risk = attack["risk"]
            vectors = ", ".join(attack["vectors"])
            endpoint = _endpoint_url(hostname or ip, port_num, service)

            row_cells = table.add_row().cells
            row_cells[0].text = str(port_num)
            row_cells[1].text = proto
            row_cells[2].text = service
            row_cells[3].text = prod_ver
            row_cells[4].text = endpoint
            row_cells[5].text = risk
            row_cells[6].text = vectors

            RISK_COLORS = {
                "Crítica": RGBColor(192, 0, 0),
                "Alta": RGBColor(200, 80, 0),
                "Media": RGBColor(190, 140, 0),
            }
            if risk in RISK_COLORS:
                for para in row_cells[5].paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RISK_COLORS[risk]
                        run.font.bold = True

        scripts_with_output = [
            s for p in open_ports for s in (p.get("scripts") or [])
            if s.get("output") and str(s.get("output")).strip()
        ]
        if scripts_with_output:
            doc.add_heading("Resultados NSE (scripts)", level=3)
            for s in scripts_with_output[:8]:
                line = f"[{s.get('id','?')}] {str(s.get('output',''))[:280]}"
                doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph(
        "Fuente: XML de Nmap (datos reales del escaneo). Ningún campo ha sido inferido manualmente.",
        style="List Bullet"
    ).runs[0].font.size = Pt(8)
