from collections import Counter, defaultdict
from datetime import datetime
import json
import os
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


# ── Colour palette ───────────────────────────────────────────────────
SEV_BG = {
    "Crítica":    "7B0000",   # dark red
    "Alta":       "C0392B",   # red
    "Media":      "D35400",   # orange
    "Baja":       "2E86C1",   # blue
    "Informativa":"616A6B",   # grey
}
SEV_FG = {
    "Crítica":    "FFFFFF",
    "Alta":       "FFFFFF",
    "Media":      "FFFFFF",
    "Baja":       "FFFFFF",
    "Informativa":"FFFFFF",
}
HEADER_BG = "1C2833"   # dark slate
HEADER_FG = "F2F3F4"


def _set_cell_bg(cell, hex_color):
    """Fill a table cell background with a hex colour (no #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_font(cell, hex_color, bold=False, size_pt=None):
    for para in cell.paragraphs:
        for run in para.runs:
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)
            if bold:
                run.font.bold = True
            if size_pt:
                run.font.size = Pt(size_pt)


def style_header_row(row, bg=HEADER_BG, fg=HEADER_FG):
    """Apply dark header style to all cells in a row."""
    for cell in row.cells:
        _set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                r = int(fg[0:2], 16)
                g = int(fg[2:4], 16)
                b = int(fg[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
                run.font.bold = True


def style_severity_cell(cell, severity):
    bg = SEV_BG.get(severity, "616A6B")
    fg = SEV_FG.get(severity, "FFFFFF")
    _set_cell_bg(cell, bg)
    _set_cell_font(cell, fg, bold=True)


SEVERITY_ORDER = {
    "Crítica": 1,
    "Alta": 2,
    "Media": 3,
    "Baja": 4,
    "Informativa": 5,
}

FINDING_STATUSES = {
    "Hallazgo",
    "Posible hallazgo",
}

ERROR_STATUSES = {
    "Error",
}

OK_STATUSES = {
    "Correcto",
    "Detectado",
    "No evidenciado",
    "No detectado",
}

OFFENSIVE_MODULES = {
    "XSS reflejado",
    "SQL Injection",
    "SQL Injection Auth (Browser)",
    "Open Redirect",
    "XSS DOM",
    "SSTI",
    "SSRF",
    "Path Traversal",
}

MODULE_PHASE = {
    "Autenticación": "Acceso inicial",
    "Enumeración de usuarios": "Enumeración",
    "Crawler": "Reconocimiento",
    "Discovery": "Reconocimiento",
    "Discovery post-login": "Post-login Discovery",
    "Mapa de URLs": "Reconocimiento",
    "Reconocimiento": "Reconocimiento",
    "Red e infraestructura": "Reconocimiento",
    "Puertos y servicios": "Reconocimiento",
    "Correlación de vulnerabilidades": "Reconocimiento",
    "Nmap reconnaissance": "Reconocimiento",
    "Nessus/Tenable": "Reconocimiento",
    "Correlación IA ofensiva": "Reconocimiento",
    "Fingerprinting avanzado": "Reconocimiento",
    "Cabeceras de seguridad": "Reconocimiento",
    "Cookies": "Reconocimiento",
    "CORS": "Reconocimiento",
    "Métodos HTTP": "Reconocimiento",
    "API Discovery": "Enumeración",
    "Formularios": "Enumeración",
    "CSRF": "Explotación",
    "XSS reflejado": "Explotación",
    "SQL Injection": "Explotación",
    "SQL Injection Auth (Browser)": "Explotación",
    "Open Redirect": "Explotación",
    "XSS DOM": "Explotación",
    "SSTI": "Explotación",
    "SSRF": "Explotación",
    "Path Traversal": "Explotación",
    "Control de acceso": "Post-explotación",
    "JWT": "Post-explotación",
    "Exposición de dependencias": "Post-explotación",
    "Aseguramiento ofensivo": "Aseguramiento",
}

PHASE_ORDER = [
    "Reconocimiento",
    "Enumeración",
    "Acceso inicial",
    "Post-login Discovery",
    "Explotación",
    "Post-explotación",
    "Aseguramiento",
    "Otros",
]

REPORTABLE_PAGE_CLASSES = {
    "auth",
    "registration",
    "protected",
    "protected_redirect_to_auth",
    "admin_candidate",
    "api_candidate",
    "sensitive_candidate",
    "server_error",
    "error_disclosure_candidate",
}

BRIEF_MAX_FINDINGS = 12
BRIEF_MAX_DISCOVERY_ROWS = 10
BRIEF_MAX_AUTH_ROWS = 6
BRIEF_MAX_ERROR_ROWS = 8

REPORT_NOISE_MODULES = {
    "Control de calidad AI",
}

IP_CONTEXT_ALLOWLIST = {
    "red e infraestructura": "Reconocimiento de red",
    "puertos y servicios": "Puertos/servicios",
    "nmap reconnaissance": "Nmap",
    "nessus/tenable": "Nessus/Tenable",
}


def safe_text(value):
    return str(value if value is not None else "")


def safe_status(value):
    try:
        return int(value)
    except Exception:
        return 0


def truncate(text, limit=450):
    text = safe_text(text)
    if len(text) <= limit:
        return text
    return text[:limit].rstrip() + "..."


def _append_unique_text(target_list: list, seen: set, value: str, *, limit: int = 220) -> bool:
    normalized = re.sub(r"\s+", " ", safe_text(value)).strip()
    if not normalized:
        return False
    key = normalized.lower()
    if key in seen:
        return False
    seen.add(key)
    target_list.append(truncate(normalized, limit))
    return True


def clean_evidence_for_report(text):
    value = safe_text(text)
    markers = [
        "| Control anti-FP: hallazgo confirmado conservado (sin descarte automático).",
        "| Control anti-FP: hallazgo confirmado conservado (sin descarte autom",
        "| FP-RISK:ALTA (evidencia débil o aislada).",
        "| FP-RISK:MEDIA (requiere corroboración adicional).",
        "| STRICT-REVIEW:PENDIENTE (doble corroboración recomendada).",
    ]
    for marker in markers:
        value = value.replace(marker, "")
    return value.strip()


def is_finding(item):
    return item.get("Resultado") in FINDING_STATUSES


def is_error(item):
    return item.get("Resultado") in ERROR_STATUSES


def is_ok(item):
    return item.get("Resultado") in OK_STATUSES


# ── Offensive coverage state ──────────────────────────────────────────
# 3-state model: EXPLOTADO / PROTEGIDO / NO CUBIERTO
EXPLOIT_RESULT = {"Hallazgo", "Posible hallazgo"}
PROTECTED_RESULT = {"No evidenciado", "No detectado", "Correcto"}
UNCOVERED_RESULT = {"Error"}

COVERAGE_BADGE = {
    "EXPLOTADO":   ("B03A2E", "FFFFFF"),   # dark red bg, white text
    "POSIBLE":     ("CA6F1E", "FFFFFF"),   # orange bg
    "PROTEGIDO":   ("1E8449", "FFFFFF"),   # green bg
    "NO CUBIERTO": ("7D6608", "FFFFFF"),   # amber bg
}


def coverage_state(resultado: str) -> str:
    if resultado == "Hallazgo":
        return "EXPLOTADO"
    if resultado == "Posible hallazgo":
        return "POSIBLE"
    if resultado in PROTECTED_RESULT:
        return "PROTEGIDO"
    return "NO CUBIERTO"


# ── Regexes for asset mining ──────────────────────────────────────────
_RE_IP = re.compile(
    r'\b(?!0\.)(?!127\.)(?!255\.)'
    r'((?:\d{1,3}\.){3}\d{1,3})\b'
)
_RE_PORT_LIST = re.compile(r'[Aa]biertos?[:\s]+([0-9][0-9, ]+)', re.I)
_RE_BANNER = re.compile(r'[Bb]anners?:([^\|]+)')
_RE_EMAIL = re.compile(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}')
_RE_VERSION = re.compile(
    r'(Apache|Nginx|PHP|jQuery|WordPress|Drupal|Tomcat|OpenSSL|Bootstrap|'
    r'Django|Laravel|Spring|Express|IIS|Joomla|Magento)[/\s]+([\d.]+)',
    re.I,
)
_RE_CRED_HINT = re.compile(
    r'(?:password|contraseña|passwd|clave|credencial)[=:\s]+(\S{3,})',
    re.I,
)
_RE_JWT = re.compile(r'eyJ[A-Za-z0-9_\-]{10,}\.[A-Za-z0-9_\-]{5,}')
_RE_SENSITIVE_PATH = re.compile(
    r'((?:/[a-zA-Z0-9_\-]{1,60}){1,6}'
    r'(?:admin|backup|config|debug|\.env|\.git|wp-admin|phpmyadmin|'
    r'manager|console|dashboard|actuator|swagger|api-docs|graphql)'
    r'[^\s|]*)',
    re.I,
)


def extract_sensitive_assets(results: list, pages: list) -> dict:
    """Mine all result evidence for concrete pentesting data."""
    assets = {
        "ips": [],
        "ip_context": {},    # ip -> set(source modules/contexts)
        "ports": [],          # list of {"port", "service", "severity", "banner"}
        "databases": [],      # list of {"type", "host", "port", "version", "source"}
        "technologies": [],   # list of {"tech", "version", "source"}
        "users": [],          # emails / usernames found
        "credentials": [],    # any credential evidence
        "jwts": [],           # JWT tokens spotted
        "sensitive_paths": [], # sensitive paths/endpoints discovered
        "api_endpoints": [],  # API routes found
        "auth_result": None,  # {"result", "url", "evidence"}
        "exposed_headers": [], # server header leakage
    }
    seen_ips = set()
    seen_tech = set()
    seen_users = set()
    seen_paths = set()
    seen_creds = set()
    seen_jwts = set()
    seen_api_endpoints = set()
    seen_headers = set()
    seen_databases = set()

    host_re = re.compile(r"host\s*:\s*([^|\s]+)", re.I)
    service_re = re.compile(r"servicio\s*:\s*([^|]+)", re.I)

    for item in results:
        module = safe_text(item.get("Módulo", "")).lower()
        evidence = safe_text(item.get("Evidencia", ""))
        desc = safe_text(item.get("Descripción", ""))
        control = safe_text(item.get("Control", ""))
        resultado = safe_text(item.get("Resultado", ""))
        text = evidence + " " + desc

        # ── IPs ───────────────────────────────────────────────────────
        for ip in _RE_IP.findall(text):
            if ip not in seen_ips:
                seen_ips.add(ip)
                assets["ips"].append(ip)
            module_name = safe_text(item.get("Módulo", "")).strip().lower()
            source_label = IP_CONTEXT_ALLOWLIST.get(module_name)
            if source_label:
                assets.setdefault("ip_context", {}).setdefault(ip, set()).add(source_label)

        # ── Open ports ────────────────────────────────────────────────
        if "puertos" in module or "port" in module or "red" in module:
            port_m = _RE_PORT_LIST.search(evidence)
            if port_m:
                port_nums = [int(p) for p in re.findall(r'\d+', port_m.group(1))]
                banners_raw = _RE_BANNER.search(evidence)
                banner_map: dict = {}
                if banners_raw:
                    for chunk in banners_raw.group(1).split(";"):
                        chunk = chunk.strip()
                        if ":" in chunk:
                            p_str, b_val = chunk.split(":", 1)
                            try:
                                banner_map[int(p_str.strip())] = b_val.strip()
                            except ValueError:
                                pass
                for p in port_nums:
                    # Severity: re-derive from known risky port ranges
                    if p in {23, 3389, 5900, 6379, 9200, 27017, 2375, 5985, 5986}:
                        sev = "Alta"
                    elif p in {21, 25, 110, 139, 445, 1433, 1521, 3306, 5432, 8080, 1883, 2181, 7001, 7002}:
                        sev = "Alta"
                    elif p in {22, 53, 111, 135, 143, 389, 636, 587, 993, 995}:
                        sev = "Media"
                    else:
                        sev = "Baja"
                    host_hint = ""
                    host_match = host_re.search(evidence)
                    if host_match:
                        host_hint = host_match.group(1).strip()

                    service_hint = ""
                    service_match = service_re.search(evidence)
                    if service_match:
                        service_hint = service_match.group(1).strip()

                    assets["ports"].append({
                        "port": p,
                        "host": host_hint,
                        "service": service_hint,
                        "banner": banner_map.get(p, ""),
                        "severity": sev,
                    })

        # ── Technologies/versions ─────────────────────────────────────
        for m in _RE_VERSION.finditer(text):
            key = (m.group(1).lower(), m.group(2))
            if key not in seen_tech:
                seen_tech.add(key)
                assets["technologies"].append({
                    "tech": m.group(1),
                    "version": m.group(2),
                    "source": module,
                })

        # ── Emails/users ──────────────────────────────────────────────
        for email in _RE_EMAIL.findall(text):
            if email.lower() not in seen_users:
                seen_users.add(email.lower())
                assets["users"].append(email)

        # ── Credentials ───────────────────────────────────────────────
        for m in _RE_CRED_HINT.finditer(text):
            _append_unique_text(assets["credentials"], seen_creds, f"{control}: {m.group(0)[:120]}", limit=160)

        # ── JWT tokens ───────────────────────────────────────────────
        for jwt in _RE_JWT.findall(text):
            _append_unique_text(assets["jwts"], seen_jwts, f"{control}: {jwt[:80]}...", limit=120)

        # ── Auth result ───────────────────────────────────────────────
        if "autent" in module or module == "autenticación":
            assets["auth_result"] = {
                "result": resultado,
                "evidence": truncate(evidence, 300),
            }

        # ── Exposed server headers ───────────────────────────────────
        if any(h in evidence.lower() for h in ["server:", "x-powered-by:", "x-aspnet"]):
            _append_unique_text(assets["exposed_headers"], seen_headers, evidence, limit=200)

        # ── Sensitive paths ───────────────────────────────────────────
        for path in _RE_SENSITIVE_PATH.findall(text):
            p = path.strip()
            if p not in seen_paths and len(p) > 3:
                seen_paths.add(p)
                assets["sensitive_paths"].append(p)

        # ── API endpoints ────────────────────────────────────────────
        if "api" in module and resultado in FINDING_STATUSES:
            _append_unique_text(assets["api_endpoints"], seen_api_endpoints, evidence, limit=260)

        # ── Database exposures from deep infra findings ───────────────
        if "base de datos expuesta detectada" in control.lower() or "mysql" in text.lower() or "postgres" in text.lower() or "mongodb" in text.lower() or "redis" in text.lower():
            db_type = ""
            db_match = re.search(r"Base de datos expuesta detectada:\s*([^|]+)", control, re.IGNORECASE)
            if db_match:
                db_type = db_match.group(1).strip()
            if not db_type:
                for name in ["MySQL/MariaDB", "PostgreSQL", "Microsoft SQL Server", "Oracle Database", "MongoDB", "Redis", "Elasticsearch", "Cassandra", "CouchDB", "Memcached"]:
                    if name.lower() in text.lower():
                        db_type = name
                        break

            host = ""
            host_match = re.search(r"Host(?:/IP)?\s*:\s*([^|]+)", evidence, re.IGNORECASE)
            if host_match:
                host = host_match.group(1).strip()

            port = ""
            port_match = re.search(r"Puerto\s*:\s*(\d+)", evidence, re.IGNORECASE)
            if port_match:
                port = port_match.group(1)

            version = ""
            version_match = re.search(r"Versión\s*:\s*([^|]+)", evidence, re.IGNORECASE)
            if version_match:
                version = version_match.group(1).strip()

            if db_type:
                key = (db_type.lower(), host.lower(), str(port).strip(), version.lower())
                if key not in seen_databases:
                    seen_databases.add(key)
                    assets["databases"].append({
                        "type": db_type,
                        "host": host,
                        "port": port,
                        "version": version,
                        "source": module,
                    })

    # Pages: extract sensitive paths from classified URLs
    for page in pages or []:
        url = safe_text(page.get("url", ""))
        classification = safe_text(page.get("classification", ""))
        if classification in {"admin_candidate", "sensitive_candidate", "api_candidate"}:
            for path in _RE_SENSITIVE_PATH.findall(url):
                p = path.strip()
                if p not in seen_paths:
                    seen_paths.add(p)
                    assets["sensitive_paths"].append(p)

    # Convert context sets to sorted lists for deterministic output.
    assets["ip_context"] = {
        ip: sorted(list(ctxs))
        for ip, ctxs in (assets.get("ip_context") or {}).items()
    }

    return assets


def is_reportable_page(page):
    status = safe_text(page.get("status_code"))
    classification = safe_text(page.get("classification"))

    if status == "404":
        return False

    if classification in ["soft_404", "request_error", "html_candidate"]:
        return False

    if classification in REPORTABLE_PAGE_CLASSES:
        return True

    if status.startswith("5"):
        return True

    return False


def get_form_detection_summary(page):
    forms = page.get("forms") or []
    browser_runtime = page.get("browser_runtime") or {}
    inputs = browser_runtime.get("inputs") or page.get("browser_inputs") or []
    buttons = browser_runtime.get("buttons") or page.get("browser_buttons") or []

    has_runtime_auth_form = any(
        "password" in str(field).lower() or "contraseña" in str(field).lower()
        for field in inputs
    ) and any(
        "email" in str(field).lower()
        or "correo" in str(field).lower()
        or "user" in str(field).lower()
        or "usuario" in str(field).lower()
        or "login" in str(field).lower()
        for field in inputs
    )

    has_browser_form = any(
        isinstance(form, dict)
        and (
            str(form.get("source", "")) == "browser_runtime"
            or str(form.get("type", "")) == "client_side_auth_form"
            or str(form.get("method", "")).lower() == "client-side/js"
        )
        for form in forms
    )

    html_forms = [
        form for form in forms
        if not (
            isinstance(form, dict)
            and (
                str(form.get("source", "")) == "browser_runtime"
                or str(form.get("type", "")) == "client_side_auth_form"
                or str(form.get("method", "")).lower() == "client-side/js"
            )
        )
    ]

    if has_runtime_auth_form or has_browser_form:
        form_type = "Formulario dinámico detectado con Playwright"
    elif html_forms:
        form_type = "Formulario HTML clásico"
    elif inputs:
        form_type = "Inputs renderizados sin clasificar"
    else:
        form_type = "No detectado"

    return {
        "forms_count": max(len(forms), 1 if (has_runtime_auth_form or has_browser_form) else 0),
        "form_type": form_type,
        "inputs_count": len(inputs),
        "buttons_count": len(buttons),
        "has_runtime_auth_form": has_runtime_auth_form,
        "has_browser_form": has_browser_form,
        "has_html_form": bool(html_forms),
    }


def get_page_observation(page):
    classification = safe_text(page.get("classification", "sin clasificar"))
    form_summary = get_form_detection_summary(page)
    forms_count = form_summary["forms_count"]

    if classification in ["auth", "registration"]:
        if forms_count:
            return f"Ruta compatible con autenticación/registro. {form_summary['form_type']}."
        return "Ruta compatible con autenticación/registro sin formulario renderizado; revisar JavaScript y endpoints API."

    if classification == "protected_redirect_to_auth":
        return "Ruta sensible que redirige a autenticación. Revisar control de acceso tras login."

    if classification == "protected":
        return "Ruta protegida detectada. Revisar autorización y exposición."

    if classification == "admin_candidate":
        if forms_count:
            return (
                f"Ruta administrativa candidata. Redirige o carga formulario de autenticación ({form_summary['form_type']}). "
                "El formulario puede pertenecer a la página de destino del redirect, no al propio recurso administrativo."
            )
        return "Ruta administrativa candidata a revisión de control de acceso."

    if classification == "api_candidate":
        return "Ruta candidata a API; revisar métodos, autenticación y errores."

    if classification == "sensitive_candidate":
        return "Ruta sensible candidata; revisar exposición de secretos, backups o configuración."

    if classification == "server_error":
        return "Error 5xx detectado; revisar filtrado de información técnica."

    if classification == "error_disclosure_candidate":
        return "Respuesta con indicadores de error técnico o posible disclosure."

    return "Elemento relevante identificado para revisión manual."


def sort_results(results):
    return sorted(
        results,
        key=lambda r: (
            SEVERITY_ORDER.get(r.get("Severidad", ""), 99),
            r.get("Módulo", ""),
            r.get("Control", ""),
        ),
    )


def _business_likelihood(item: dict) -> str:
    status = safe_text(item.get("Resultado", "")).lower()
    conf = item.get("Confianza", 0)
    try:
        conf = float(conf or 0)
    except Exception:
        conf = 0.0
    if status == "hallazgo" or conf >= 0.8:
        return "Alta"
    if status == "posible hallazgo" or conf >= 0.55:
        return "Media"
    return "Baja"


def _business_impact(item: dict) -> str:
    sev = safe_text(item.get("Severidad", ""))
    if sev in {"Crítica", "Alta"}:
        return "Alto"
    if sev == "Media":
        return "Medio"
    return "Bajo"


def _business_risk(probability: str, impact: str) -> str:
    matrix = {
        ("Alta", "Alto"): "Crítico",
        ("Alta", "Medio"): "Alto",
        ("Alta", "Bajo"): "Medio",
        ("Media", "Alto"): "Alto",
        ("Media", "Medio"): "Medio",
        ("Media", "Bajo"): "Bajo",
        ("Baja", "Alto"): "Medio",
        ("Baja", "Medio"): "Bajo",
        ("Baja", "Bajo"): "Bajo",
    }
    return matrix.get((probability, impact), "Bajo")


def add_business_risk_matrix(document, findings):
    document.add_heading("5.1 Matriz de riesgo de negocio (Probabilidad x Impacto)", 2)

    if not findings:
        document.add_paragraph("Sin hallazgos priorizados para construir matriz de riesgo de negocio.")
        return

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Control", "Módulo", "Sev.", "Probabilidad", "Impacto", "Riesgo negocio"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    style_header_row(table.rows[0])

    for item in _prioritized_findings(findings, limit=10):
        probability = _business_likelihood(item)
        impact = _business_impact(item)
        risk = _business_risk(probability, impact)

        row = table.add_row().cells
        row[0].text = truncate(item.get("Control", ""), 95)
        row[1].text = truncate(item.get("Módulo", ""), 45)
        row[2].text = safe_text(item.get("Severidad", ""))
        row[3].text = probability
        row[4].text = impact
        row[5].text = risk


def add_compliance_mapping(document, findings):
    document.add_heading("5.2 Mapeo de cumplimiento (OWASP/NIST/ISO/GDPR)", 2)

    if not findings:
        document.add_paragraph("Sin hallazgos para mapear a marcos de cumplimiento.")
        return

    mappings = []
    seen = set()
    for item in findings:
        module = safe_text(item.get("Módulo", ""))
        control = safe_text(item.get("Control", "")).lower()
        sev = safe_text(item.get("Severidad", ""))

        refs = []
        if any(k in control for k in ["strict-transport-security", "hsts", "ftp", "tls", "http alternativo"]):
            refs.append("GDPR Art.32 (confidencialidad/cifrado)")
            refs.append("ISO27001 A.8.24 / A.8.25")
            refs.append("NIST PR.DS-2")
        if any(k in control for k in ["content-security-policy", "x-frame-options", "x-content-type-options", "cabecera"]):
            refs.append("OWASP ASVS V14 (configuración segura)")
            refs.append("NIST PR.IP-1")
        if module in {"Control de acceso", "JWT", "Autenticación"}:
            refs.append("OWASP API Top 10 - API1/API2")
            refs.append("ISO27001 A.5.15 / A.8.2")
        if module in {"SQL Injection", "XSS reflejado", "SSTI", "Path Traversal", "SSRF"}:
            refs.append("OWASP Top 10 A03/A05")
            refs.append("NIST SI-10")

        if not refs:
            continue

        key = (module, control)
        if key in seen:
            continue
        seen.add(key)
        mappings.append({
            "control": truncate(item.get("Control", ""), 120),
            "module": module,
            "severity": sev,
            "refs": "; ".join(sorted(set(refs))),
        })

    if not mappings:
        document.add_paragraph("No se generaron mapeos normativos automáticos en esta ejecución.")
        return

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Control", "Módulo", "Sev.", "Referencias normativas"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    style_header_row(table.rows[0])

    for row_data in mappings[:12]:
        row = table.add_row().cells
        row[0].text = row_data["control"]
        row[1].text = row_data["module"]
        row[2].text = row_data["severity"]
        row[3].text = row_data["refs"]


def add_manual_offensive_backlog(document, findings, pages, assets):
    document.add_heading("5.3 Backlog ofensivo manual pendiente", 2)
    document.add_paragraph(
        "Pruebas manuales prioritarias para cerrar brecha entre reconocimiento automatizado y auditoría ofensiva completa."
    )

    backlog = []
    auth_pages = [p for p in pages or [] if safe_text(p.get("classification", "")).lower() in {"auth", "registration", "protected", "admin_candidate"}]
    if auth_pages:
        backlog.append("Validar BOLA/IDOR por rol en rutas post-login (acceso cruzado entre entidades/usuarios).")
        backlog.append("Validar manipulación de sesión/JWT (claims, expiración, firma, replay y privilegios).")

    sensitive_forms = any(get_form_detection_summary(p).get("forms_count", 0) > 0 for p in (pages or []))
    if sensitive_forms:
        backlog.append("Ejecutar pruebas manuales en formularios dinámicos (SQLi/XSS/CSRF) con variaciones de contexto.")

    ports = {int(p.get("port", 0) or 0) for p in (assets.get("ports") or []) if str(p.get("port", "")).isdigit()}
    if 21 in ports:
        backlog.append("Comprobar manualmente FTP: acceso anónimo, cifrado real, exposición de credenciales y hardening de servicio.")
    if 8080 in ports:
        backlog.append("Auditar manualmente servicio en 8080: paneles admin, APIs no documentadas y controles de autenticación.")

    if not backlog:
        backlog.append("No se detectaron prerrequisitos claros de explotación manual; revisar alcance y profundidad de crawling autenticado.")

    for item in backlog[:8]:
        document.add_paragraph(item, style="List Bullet")


def add_defensive_playbook(document, assets, findings):
    document.add_heading("5.4 Playbook defensivo técnico (acción inmediata)", 2)

    techs = {safe_text(t.get("tech", "")).lower() for t in (assets.get("technologies") or [])}
    controls = " ".join(safe_text(f.get("Control", "")).lower() for f in (findings or []))

    has_header_gaps = any(k in controls for k in [
        "content-security-policy",
        "strict-transport-security",
        "x-content-type-options",
        "x-frame-options",
    ])

    if has_header_gaps or "nginx" in techs:
        document.add_paragraph("Nginx hardening recomendado:", style="List Bullet")
        document.add_paragraph(
            "add_header Strict-Transport-Security \"max-age=31536000; includeSubDomains; preload\" always;\n"
            "add_header X-Frame-Options \"SAMEORIGIN\" always;\n"
            "add_header X-Content-Type-Options \"nosniff\" always;\n"
            "add_header Referrer-Policy \"strict-origin-when-cross-origin\" always;\n"
            "server_tokens off;"
        )

    if has_header_gaps or "nextjs" in techs or "next.js" in techs:
        document.add_paragraph("Next.js CSP base sugerida (ajustar dominios permitidos):", style="List Bullet")
        document.add_paragraph(
            "Content-Security-Policy: default-src 'self'; script-src 'self' 'unsafe-inline' https:; "
            "style-src 'self' 'unsafe-inline'; img-src 'self' data: https:; connect-src 'self' https:; "
            "font-src 'self' data:; frame-ancestors 'none'; base-uri 'self';"
        )

    document.add_paragraph("Defensa en capas recomendada:", style="List Bullet")
    document.add_paragraph(
        "Aplicar rate-limiting en /login y /registro, WAF/CDN con reglas anti-bot y geoblocking según exposición, "
        "y segmentación de puertos administrativos fuera de Internet pública."
    )


def _version_anomaly_notes(technologies: list) -> list:
    notes = []
    for tech in technologies or []:
        name = safe_text(tech.get("tech")).strip().lower()
        version = safe_text(tech.get("version")).strip()
        if not name or not version:
            continue

        parts = version.split(".")
        try:
            major = int(parts[0]) if parts else 0
            minor = int(parts[1]) if len(parts) > 1 else 0
        except Exception:
            continue

        # Conservative heuristic: alert when banner version looks implausibly ahead.
        if name == "nginx" and (major > 1 or (major == 1 and minor >= 29)):
            notes.append(
                f"Versión reportada potencialmente anómala para {name}/{version}: validar posible banner spoofing, proxy intermedio o identificación errónea."
            )

    return notes


def _extract_url_hint(text: str) -> str:
    value = safe_text(text)

    tagged_patterns = [
        r"Final:\s*(https?://[^\s|]+)",
        r"Solicitada:\s*(https?://[^\s|]+)",
        r"URL final:\s*(https?://[^\s|]+)",
        r"Ruta post-login:\s*(https?://[^\s|]+)",
    ]
    for pattern in tagged_patterns:
        m = re.search(pattern, value, re.I)
        if m:
            return m.group(1).strip().rstrip("/")

    generic = re.search(r"(https?://[^\s|]+)", value, re.I)
    if generic:
        return generic.group(1).strip().rstrip("/")

    return ""


def _normalize_control_for_key(control: str) -> str:
    text = safe_text(control).strip().lower()
    text = re.sub(r"\s+", " ", text)

    if text.startswith("ruta descubierta -"):
        return "ruta_descubierta"
    if "discovery de superficie web" in text:
        return "discovery_superficie"
    if "cobertura post-login" in text:
        return "cobertura_post_login"

    if "content-security-policy" in text or "csp" in text:
        return "missing_csp"
    if "strict-transport-security" in text or "hsts" in text:
        return "missing_hsts"
    if "x-content-type-options" in text:
        return "missing_x_content_type_options"
    if "x-frame-options" in text:
        return "missing_x_frame_options"
    if "servicios con riesgo de exposición" in text or "puerto crítico expuesto" in text:
        return "network_port_exposure"

    return text


def _finding_cluster_key(item: dict) -> tuple:
    module = safe_text(item.get("Módulo", "")).strip().lower()
    control_norm = _normalize_control_for_key(item.get("Control", ""))

    # Collapse duplicated header findings across modules (headers/correlation).
    if control_norm in {
        "missing_csp",
        "missing_hsts",
        "missing_x_content_type_options",
        "missing_x_frame_options",
    }:
        return ("header_gap", control_norm)

    # Collapse port exposure duplicates across infra modules.
    if control_norm == "network_port_exposure" or module in {"puertos y servicios", "nmap reconnaissance"}:
        if "puerto crítico expuesto" in safe_text(item.get("Control", "")).lower():
            return ("network_port_exposure", safe_text(item.get("Control", "")).strip().lower())
        return ("network_port_exposure", "general")

    return (module, control_norm)


def _result_dedupe_key(item: dict) -> tuple:
    if not isinstance(item, dict):
        value = safe_text(item).strip().lower()
        return ("raw", value, "", "")

    module = safe_text(item.get("Módulo", "")).strip().lower()
    control = _normalize_control_for_key(item.get("Control", ""))
    evidence = safe_text(item.get("Evidencia", ""))
    url_hint = _extract_url_hint(evidence)

    if control == "ruta_descubierta":
        return (control, url_hint)

    return (module, control, url_hint)


def _extract_int_metric(evidence: str, label: str) -> int:
    m = re.search(rf"{re.escape(label)}\s*:\s*(\d+)", safe_text(evidence), re.I)
    if not m:
        return 0
    try:
        return int(m.group(1))
    except Exception:
        return 0


def _is_low_value_repeated_row(item: dict) -> bool:
    module = safe_text(item.get("Módulo", "")).strip()
    control = safe_text(item.get("Control", "")).strip().lower()
    status = safe_text(item.get("Resultado", "")).strip()
    evidence = safe_text(item.get("Evidencia", ""))

    if module in REPORT_NOISE_MODULES:
        return True

    # Ignore repetitive post-login summaries when there was no net change.
    if module == "Autenticación" and "cobertura post-login" in control and status == "Detectado":
        post_urls = _extract_int_metric(evidence, "URLs post-login")
        protected = _extract_int_metric(evidence, "Rutas protegidas detectadas")
        if post_urls <= 0 and protected <= 0:
            return True

    # Remove discovery noise from known irrelevant responses.
    if module in {"Discovery", "Discovery post-login"}:
        blob = f"{control} {evidence}".lower()
        if any(tok in blob for tok in ["404", "soft_404", "html_candidate", "request_error"]):
            return True

    return False


def _is_meaningful_page_for_report(page: dict) -> bool:
    status = safe_status(page.get("status_code"))
    classification = safe_text(page.get("classification", "")).lower()
    final_url = safe_text(page.get("final_url") or page.get("url"))

    if not final_url:
        return False
    if status == 404:
        return False
    if classification in {"soft_404", "request_error", "html_candidate"}:
        return False

    return classification in REPORTABLE_PAGE_CLASSES or status in (401, 403) or status >= 500


def _dedupe_pages_for_report(pages: list) -> list:
    best = {}
    order = []

    def _page_score(page):
        status = safe_status(page.get("status_code"))
        classification = safe_text(page.get("classification", "")).lower()
        score = 0
        if 200 <= status < 400:
            score += 4
        elif status in (401, 403):
            score += 3
        elif status >= 500:
            score += 1

        if classification in {"protected", "protected_redirect_to_auth", "admin_candidate", "api_candidate", "sensitive_candidate"}:
            score += 2
        if safe_text(page.get("discovery_context", "")).lower() == "post_login":
            score += 1
        return score

    for page in pages or []:
        key = safe_text(page.get("final_url") or page.get("url")).strip().rstrip("/")
        if not key:
            continue
        if key not in best:
            best[key] = page
            order.append(key)
            continue
        if _page_score(page) > _page_score(best[key]):
            best[key] = page

    return [best[key] for key in order]


def _prioritized_findings(findings: list, limit: int = BRIEF_MAX_FINDINGS) -> list:
    unique = {}
    order = []
    for item in sort_results(findings or []):
        key = _finding_cluster_key(item)
        current = unique.get(key)
        if not current:
            unique[key] = item
            order.append(key)
            continue

        cur_rank = (
            SEVERITY_ORDER.get(safe_text(current.get("Severidad")), 99),
            -float(current.get("Confianza", 0) or 0),
        )
        new_rank = (
            SEVERITY_ORDER.get(safe_text(item.get("Severidad")), 99),
            -float(item.get("Confianza", 0) or 0),
        )
        if new_rank < cur_rank:
            unique[key] = item

    return [unique[key] for key in order][:limit]


def dedupe_results_for_report(results):
    """Drop noisy duplicates while preserving strongest status/severity rows."""
    rank_status = {
        "Hallazgo": 5,
        "Posible hallazgo": 4,
        "Error": 3,
        "Detectado": 2,
        "Correcto": 1,
        "No evidenciado": 1,
        "No detectado": 1,
    }
    rank_sev = {
        "Crítica": 5,
        "Alta": 4,
        "Media": 3,
        "Baja": 2,
        "Informativa": 1,
    }

    best = {}
    order = []
    for item in results or []:
        if not isinstance(item, dict):
            continue

        if _is_low_value_repeated_row(item):
            continue

        module = safe_text(item.get("Módulo", "")).strip()
        status = safe_text(item.get("Resultado", "")).strip()
        sev = safe_text(item.get("Severidad", "")).strip()
        key = _result_dedupe_key(item)
        conf = item.get("Confianza", 0)
        try:
            conf_score = float(conf or 0)
        except Exception:
            conf_score = 0.0
        score = (rank_status.get(status, 0), rank_sev.get(sev, 0), conf_score)

        current = best.get(key)
        if not current or score > current[0]:
            best[key] = (score, dict(item))
            if key not in order:
                order.append(key)

    return [best[key][1] for key in order if key in best]


def set_document_style(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True


def _apply_badge(cell, badge_key: str):
    """Fill a table cell with coverage badge colour."""
    bg, fg = COVERAGE_BADGE.get(badge_key, ("616A6B", "FFFFFF"))
    _set_cell_bg(cell, bg)
    _set_cell_font(cell, fg, bold=True)


def add_sensitive_assets_section(document, assets: dict):
    """Section 2 — concrete intelligence extracted from scan results."""
    document.add_heading("2. Activos y datos sensibles descubiertos", 1)

    document.add_paragraph(
        "Resumen consolidado de exposición técnica relevante para priorización de riesgos."
    )

    # ── Authentication result ─────────────────────────────────────────
    auth = assets.get("auth_result")
    if auth:
        document.add_heading("Estado de autenticación", 3)
        result_label = auth.get("result", "Desconocido")
        badge = {
            "Autenticado": "PROTEGIDO",
            "Falló": "EXPLOTADO",    # auth bypass or valid creds worked
            "Indeterminado": "NO CUBIERTO",
        }.get(result_label, "NO CUBIERTO")
        p = document.add_paragraph()
        run = p.add_run(f"  {result_label}  ")
        bg_h, fg_h = COVERAGE_BADGE.get(badge, ("616A6B", "FFFFFF"))
        r, g, b = int(bg_h[:2], 16), int(bg_h[2:4], 16), int(bg_h[4:], 16)
        run.font.color.rgb = RGBColor(
            int(fg_h[:2], 16), int(fg_h[2:4], 16), int(fg_h[4:], 16)
        )
        run.font.bold = True
        document.add_paragraph(f"Evidencia: {truncate(auth.get('evidence', ''), 220)}")

    # ── IPs ────────────────────────────────────────────────────────────
    ips = assets.get("ips") or []
    if ips:
        document.add_heading("IPs/Hosts descubiertos", 3)
        ip_context = assets.get("ip_context") or {}
        for ip in ips[:10]:
            contexts = ip_context.get(ip) or []
            if contexts:
                document.add_paragraph(
                    f"{ip} — evidencia técnica: {', '.join(contexts[:3])}",
                    style="List Bullet",
                )
            else:
                document.add_paragraph(ip, style="List Bullet")

    # ── Open ports ────────────────────────────────────────────────────
    ports = assets.get("ports") or []
    if ports:
        document.add_heading(f"Puertos abiertos ({len(ports)} encontrados)", 3)
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0]
        for i, h in enumerate(["Host", "Puerto", "Servicio", "Severidad", "Banner", "Riesgo"]):
            hdr.cells[i].text = h
        style_header_row(hdr)
        seen_p = set()
        for pinfo in sorted(ports, key=lambda x: (
            {"Alta": 0, "Media": 1, "Baja": 2}.get(x.get("severity", "Baja"), 3),
            x.get("port", 0),
        ))[:10]:
            p_num = pinfo.get("port", 0)
            if p_num in seen_p:
                continue
            seen_p.add(p_num)
            row = table.add_row().cells
            row[0].text = pinfo.get("host", "") or "—"
            row[1].text = str(p_num)
            row[2].text = pinfo.get("service", "") or "—"
            sev = pinfo.get("severity", "Baja")
            row[3].text = sev
            row[4].text = pinfo.get("banner", "") or "—"
            row[5].text = "⚠ CRÍTICO" if p_num in {23, 3389, 5900, 6379, 9200, 27017, 2375} else "Revisar"
            style_severity_cell(row[3], sev)

    # ── Technologies ──────────────────────────────────────────────────
    techs = assets.get("technologies") or []
    if techs:
        document.add_heading("Tecnologías y versiones detectadas", 3)
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0]
        for i, h in enumerate(["Tecnología", "Versión", "Origen"]):
            hdr.cells[i].text = h
        style_header_row(hdr)
        for t in techs[:8]:
            row = table.add_row().cells
            row[0].text = t.get("tech", "")
            row[1].text = t.get("version", "")
            row[2].text = t.get("source", "")

        anomaly_notes = _version_anomaly_notes(techs)
        for note in anomaly_notes[:3]:
            document.add_paragraph(note, style="List Bullet")

    # ── Users/Emails ──────────────────────────────────────────────────
    users = assets.get("users") or []
    if users:
        document.add_heading(f"Usuarios/Emails encontrados ({len(users)})", 3)
        document.add_paragraph("Se detectaron identificadores de usuario/correo en respuestas. Revisar exposición y minimización de datos.")

    # ── Credentials ───────────────────────────────────────────────────
    creds = assets.get("credentials") or []
    if creds:
        document.add_heading(f"⚠ CREDENCIALES / DATOS SENSIBLES ({len(creds)} indicadores)", 3)
        for c in creds[:5]:
            document.add_paragraph(c, style="List Bullet")
    else:
        document.add_heading("Credenciales", 3)
        document.add_paragraph("No se encontraron credenciales expuestas en esta ejecución.")

    # ── JWT tokens ────────────────────────────────────────────────────
    jwts = assets.get("jwts") or []
    if jwts:
        document.add_heading(f"Tokens JWT detectados ({len(jwts)})", 3)
        for j in jwts[:5]:
            document.add_paragraph(j, style="List Bullet")

    # ── Sensitive paths ───────────────────────────────────────────────
    paths = assets.get("sensitive_paths") or []
    if paths:
        document.add_heading(f"Rutas/Endpoints sensibles ({len(paths)} detectados)", 3)
        for p in sorted(set(paths))[:10]:
            document.add_paragraph(p, style="List Bullet")

    # ── Databases ─────────────────────────────────────────────────────
    dbs = assets.get("databases") or []
    if dbs:
        document.add_heading(f"Bases de datos detectadas ({len(dbs)})", 3)
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["Tipo", "Host/IP", "Puerto", "Versión", "Origen"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        style_header_row(table.rows[0])

        for db in dbs[:10]:
            row = table.add_row().cells
            row[0].text = db.get("type", "")
            row[1].text = db.get("host", "") or "—"
            row[2].text = str(db.get("port", "") or "—")
            row[3].text = db.get("version", "") or "—"
            row[4].text = db.get("source", "") or "—"

    # ── API endpoints ─────────────────────────────────────────────────
    apis = assets.get("api_endpoints") or []
    if apis:
        document.add_heading(f"Endpoints API descubiertos ({len(apis)})", 3)
        for a in apis[:6]:
            document.add_paragraph(a, style="List Bullet")

    # ── Exposed server headers ────────────────────────────────────────
    hdrs = assets.get("exposed_headers") or []
    if hdrs:
        document.add_heading("Cabeceras que revelan tecnología", 3)
        for h in hdrs[:4]:
            document.add_paragraph(h, style="List Bullet")

    if not any([ips, ports, techs, users, creds, jwts, paths, apis, dbs, hdrs, auth]):
        document.add_paragraph(
            "No se extrajo inteligencia técnica concreta. "
            "Esto puede indicar que el objetivo usa ofuscación o que los módulos de reconocimiento necesitan credenciales."
        )


def add_logo(document):
    logo_path = "Logo_horizontal.png"

    if os.path.exists(logo_path):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Inches(1.65))


def add_title_page(document, audit_name, target_url, scan_mode=None, pages_count=None):
    add_logo(document)

    title = document.add_heading(
        "INFORME TÉCNICO DE AUDITORÍA DE SEGURIDAD WEB - BLACKHARRIER WEB SENTINEL",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")
    document.add_paragraph(f"Proyecto: {audit_name}")
    document.add_paragraph(f"Objetivo evaluado: {target_url}")

    if scan_mode:
        document.add_paragraph(f"Modo de auditoría: {scan_mode}")

    if pages_count is not None:
        document.add_paragraph(f"URLs HTML analizadas por crawler: {pages_count}")

    document.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    document.add_paragraph("Clasificación: Uso interno / Auditoría autorizada")


def add_summary_table(document, findings, errors, oks):
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    table.rows[0].cells[0].text = "Indicador"
    table.rows[0].cells[1].text = "Cantidad"

    rows = [
        ("Hallazgos", len(findings)),
        ("Errores de ejecución", len(errors)),
        ("Comprobaciones sin hallazgo / informativas", len(oks)),
    ]

    for name, count in rows:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = str(count)


def _extract_candidate_targets_from_results(results: list) -> list:
    candidates = []
    for row in results or []:
        control = safe_text(row.get("Control", "")).lower()
        if "objetivos candidatos priorizados" not in control:
            continue
        evidence = safe_text(row.get("Evidencia", "")).strip()
        if not evidence:
            continue
        try:
            parsed = json.loads(evidence)
        except Exception:
            continue
        if isinstance(parsed, list):
            candidates.extend([x for x in parsed if isinstance(x, dict)])
    return candidates


def _estimate_remediation_effort(findings: list, assets: dict) -> str:
    crit = len([x for x in findings if safe_text(x.get("Severidad")) == "Crítica"])
    high = len([x for x in findings if safe_text(x.get("Severidad")) == "Alta"])
    exposed_ports = len([x for x in (assets.get("ports") or []) if safe_text(x.get("severity")) in {"Alta", "Media"}])
    if crit > 0 or high >= 8 or exposed_ports >= 6:
        return "Medio-Alto"
    if high >= 3 or exposed_ports >= 3:
        return "Medio"
    return "Bajo-Medio"


def add_c_level_one_page_summary(document, target_url, findings, errors, assets, deduped_results):
    document.add_heading("1.1 Resumen Ejecutivo C-Level (1 página)", 2)

    high_count = len([x for x in findings if safe_text(x.get("Severidad")) in {"Crítica", "Alta"}])
    medium_count = len([x for x in findings if safe_text(x.get("Severidad")) == "Media"])
    risky_ports = sorted({int(p.get("port", 0) or 0) for p in (assets.get("ports") or []) if safe_text(p.get("severity")) in {"Alta", "Media"}})
    candidate_targets = _extract_candidate_targets_from_results(deduped_results)
    effort = _estimate_remediation_effort(findings, assets)

    risk_level = "Medio"
    if high_count >= 6:
        risk_level = "Alto"
    elif high_count <= 1 and medium_count <= 2:
        risk_level = "Bajo-Medio"

    document.add_paragraph(
        f"El estado de seguridad actual de {target_url} presenta un riesgo {risk_level} para continuidad y confidencialidad, "
        "con superficie expuesta en capas de infraestructura y aplicación que requiere remediación priorizada."
    )

    kpi = document.add_table(rows=1, cols=5)
    kpi.style = "Table Grid"
    headers = ["Riesgo global", "Hallazgos Altos/Críticos", "Hallazgos Medios", "Errores ejecución", "Esfuerzo remediación"]
    for i, h in enumerate(headers):
        kpi.rows[0].cells[i].text = h
    style_header_row(kpi.rows[0])

    row = kpi.add_row().cells
    row[0].text = risk_level
    row[1].text = str(high_count)
    row[2].text = str(medium_count)
    row[3].text = str(len(errors))
    row[4].text = effort

    document.add_paragraph("Impacto de negocio potencial:")
    document.add_paragraph(
        "Exposición de servicios y configuraciones de seguridad insuficientes pueden habilitar fuga de credenciales, "
        "interrupción del servicio y aumento de superficie de ataque.",
        style="List Bullet",
    )
    if risky_ports:
        document.add_paragraph(
            f"Puertos con riesgo operativo observado: {', '.join(str(p) for p in risky_ports[:8])}.",
            style="List Bullet",
        )
    if candidate_targets:
        top = sorted(candidate_targets, key=lambda x: float(x.get("priority_score", 0) or 0), reverse=True)[:3]
        compact = " | ".join(
            f"{safe_text(x.get('target'))} (score {safe_text(x.get('priority_score'))})"
            for x in top
        )
        document.add_paragraph(
            f"Activos prioritarios para mitigación inmediata: {compact}",
            style="List Bullet",
        )

    document.add_paragraph("Decisiones ejecutivas recomendadas (30-60 días):")
    document.add_paragraph("Bloquear exposición innecesaria de servicios y paneles administrativos desde Internet.", style="List Number")
    document.add_paragraph("Aplicar baseline obligatorio de cabeceras y cifrado TLS en todos los frontales web.", style="List Number")
    document.add_paragraph("Completar validación ofensiva manual sobre autenticación, sesión y control de acceso por roles.", style="List Number")


def _asset_row_score(base_score: float, sev: str, kind: str) -> float:
    sev_bonus = {
        "Crítica": 35,
        "Alta": 25,
        "Media": 15,
        "Baja": 8,
        "Informativa": 3,
    }.get(sev, 5)
    kind_bonus = {
        "port": 20,
        "endpoint": 16,
        "candidate": 14,
        "identity": 10,
    }.get(kind, 8)
    return max(0.0, min(base_score + sev_bonus + kind_bonus, 100.0))


def build_asset_risk_register(findings: list, assets: dict, results: list) -> list:
    rows = []

    for p in assets.get("ports") or []:
        port = int(p.get("port", 0) or 0)
        sev = "Alta" if safe_text(p.get("severity")) == "Alta" else "Media"
        score = _asset_row_score(35, sev, "port")
        rows.append({
            "asset": f"Port {port}",
            "kind": "Servicio de red",
            "severity": sev,
            "score": round(score, 1),
            "reason": safe_text(p.get("banner", "")) or "Puerto expuesto con necesidad de hardening.",
            "action": "Restringir exposición por firewall/ACL, validar autenticación y parcheo.",
        })

    for ep in (assets.get("api_endpoints") or [])[:10]:
        score = _asset_row_score(32, "Media", "endpoint")
        rows.append({
            "asset": truncate(ep, 90),
            "kind": "Endpoint/API",
            "severity": "Media",
            "score": round(score, 1),
            "reason": "Endpoint descubierto con potencial exposición funcional.",
            "action": "Validar authN/authZ por recurso, rate-limit y validación de entrada.",
        })

    for path in (assets.get("sensitive_paths") or [])[:10]:
        score = _asset_row_score(38, "Media", "endpoint")
        rows.append({
            "asset": truncate(path, 90),
            "kind": "Ruta sensible",
            "severity": "Media",
            "score": round(score, 1),
            "reason": "Ruta sensible identificada en discovery/crawling.",
            "action": "Aplicar control de acceso estricto y ocultar rutas administrativas no públicas.",
        })

    for user in (assets.get("users") or [])[:6]:
        score = _asset_row_score(28, "Baja", "identity")
        rows.append({
            "asset": user,
            "kind": "Identidad expuesta",
            "severity": "Baja",
            "score": round(score, 1),
            "reason": "Identificador de usuario/correo observado en respuestas.",
            "action": "Reducir exposición de PII y aplicar minimización de datos en respuestas.",
        })

    for c in _extract_candidate_targets_from_results(results)[:12]:
        base = float(c.get("priority_score", 0) or 0) * 7.5
        sev = "Alta" if float(c.get("priority_score", 0) or 0) >= 8 else "Media"
        score = _asset_row_score(base, sev, "candidate")
        rows.append({
            "asset": safe_text(c.get("target", "")),
            "kind": safe_text(c.get("kind", "Candidato")),
            "severity": sev,
            "score": round(score, 1),
            "reason": safe_text(c.get("reason", "")) or "Activo priorizado por correlación ofensiva.",
            "action": "Validar alcance legal y ejecutar hardening/remediación priorizada por impacto.",
        })

    # Collapse duplicates by asset+kind keeping highest score.
    dedup = {}
    for row in rows:
        key = (safe_text(row.get("asset", "")).lower(), safe_text(row.get("kind", "")).lower())
        current = dedup.get(key)
        if not current or float(row.get("score", 0) or 0) > float(current.get("score", 0) or 0):
            dedup[key] = row

    final_rows = sorted(dedup.values(), key=lambda x: float(x.get("score", 0) or 0), reverse=True)
    return final_rows[:16]


def add_asset_risk_scoring_section(document, findings, assets, results):
    document.add_heading("5.5 Priorización por activo (score 0-100)", 2)

    risk_rows = build_asset_risk_register(findings, assets, results)
    if not risk_rows:
        document.add_paragraph("No se pudo construir scoring por activo con la evidencia disponible.")
        return

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Activo", "Tipo", "Sev.", "Score", "Motivo", "Acción recomendada"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    style_header_row(table.rows[0])

    for entry in risk_rows:
        row = table.add_row().cells
        row[0].text = truncate(entry.get("asset", ""), 80)
        row[1].text = truncate(entry.get("kind", ""), 35)
        sev = safe_text(entry.get("severity", "Media"))
        row[2].text = sev
        row[3].text = str(entry.get("score", ""))
        row[4].text = truncate(entry.get("reason", ""), 140)
        row[5].text = truncate(entry.get("action", ""), 160)
        style_severity_cell(row[2], sev)


def add_severity_table(document, findings):
    document.add_heading("Resumen por severidad de hallazgos", 2)

    counts = Counter([x.get("Severidad", "Sin clasificar") for x in findings])

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    hdr = table.rows[0]
    hdr.cells[0].text = "Severidad"
    hdr.cells[1].text = "Hallazgos"
    style_header_row(hdr)

    for sev in ["Crítica", "Alta", "Media", "Baja", "Informativa"]:
        row = table.add_row().cells
        row[0].text = sev
        row[1].text = str(counts.get(sev, 0))
        style_severity_cell(row[0], sev)


def add_discovered_surface_body(document, pages):
    """Body of section 3 — discovered surface table (no section heading)."""
    pages = _dedupe_pages_for_report(pages or [])
    reportable_pages = [page for page in pages if _is_meaningful_page_for_report(page)]

    if not reportable_pages:
        document.add_paragraph(
            "No se identificaron URLs relevantes reportables. "
            "Las rutas inexistentes, 404, soft-404 y páginas genéricas se han omitido para reducir ruido."
        )
        return

    document.add_paragraph("Listado depurado de rutas con valor técnico para validación manual.")

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    headers = ["URL origen", "URL final", "HTTP", "Clasificación", "Formularios", "Observación"]
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        hdr_row.cells[i].text = h
    style_header_row(hdr_row)

    for page in reportable_pages[:BRIEF_MAX_DISCOVERY_ROWS]:
        url = safe_text(page.get("url"))
        final_url = safe_text(page.get("final_url") or url)
        status_code = safe_text(page.get("status_code"))
        classification = safe_text(page.get("classification", "sin clasificar"))
        form_summary = get_form_detection_summary(page)
        forms_count = form_summary["forms_count"]
        observation = get_page_observation(page)

        row = table.add_row().cells
        row[0].text = truncate(url, 120)
        row[1].text = truncate(final_url, 120)
        row[2].text = status_code
        row[3].text = classification
        row[4].text = str(forms_count)
        row[5].text = truncate(observation, 220)


def add_discovered_surface(document, pages):
    """Legacy wrapper — kept for compatibility."""
    document.add_heading("2. Superficie relevante descubierta por crawler/discovery", 1)
    add_discovered_surface_body(document, pages)


def add_auth_surface_summary(document, pages):
    auth_keywords = [
        "login",
        "signin",
        "auth",
        "registro",
        "register",
        "signup",
        "iniciar-sesion",
        "inicio-sesion",
        "crear-cuenta",
    ]

    auth_pages = [
        page for page in _dedupe_pages_for_report(pages or [])
        if is_reportable_page(page)
        and (
            page.get("classification") in ["auth", "registration"]
            or any(x in safe_text(page.get("url")).lower() for x in auth_keywords)
            or any(x in safe_text(page.get("final_url")).lower() for x in auth_keywords)
        )
    ]

    document.add_heading("4. Superficie de autenticación y registro", 1)

    if not auth_pages:
        document.add_paragraph(
            "No se identificaron rutas válidas de autenticación o registro en el alcance reportable."
        )
        return

    document.add_paragraph("Superficie de autenticación priorizada para pruebas de acceso, sesión y control de autorización.")

    for page in auth_pages[:BRIEF_MAX_AUTH_ROWS]:
        url = page.get("url")
        final_url = page.get("final_url") or url
        form_summary = get_form_detection_summary(page)

        document.add_paragraph(
            f"Ruta detectada: {url} | URL final: {final_url} | HTTP: {page.get('status_code')} | "
            f"Clasificación: {page.get('classification')} | "
            f"Tipo formulario: {form_summary['form_type']} | "
            f"Formularios detectados: {form_summary['forms_count']} | "
            f"Inputs renderizados: {form_summary['inputs_count']} | Botones/enlaces: {form_summary['buttons_count']}",
            style="List Bullet",
        )


def add_authenticated_session_evidence(document, results, pages):
    document.add_heading("4.1 Evidencia de sesión autenticada y expansión post-login", 2)

    auth_rows = [
        row for row in results or []
        if safe_text(row.get("Módulo")).strip() == "Autenticación"
        or "autentic" in safe_text(row.get("Control")).lower()
        or safe_text(row.get("Control")).strip() == "Cobertura post-login"
    ]

    if not auth_rows:
        document.add_paragraph(
            "No se registró evidencia estructurada de autenticación en los resultados normalizados."
        )
        return

    status_rank = {
        "Autenticado": 0,
        "Detectado": 1,
        "Indeterminado": 2,
        "Fallido": 3,
        "No configurado": 4,
        "Error": 5,
    }

    def auth_row_rank(row):
        control = safe_text(row.get("Control")).lower()
        status = safe_text(row.get("Resultado"))
        has_post_login = "post-login" in control or "cobertura post-login" in control
        has_cookie_evidence = "cookies" in safe_text(row.get("Evidencia")).lower()
        return (
            0 if has_post_login else 1,
            0 if has_cookie_evidence else 1,
            status_rank.get(status, 99),
        )

    best_auth = sorted(auth_rows, key=auth_row_rank)[0]

    document.add_paragraph(
        f"Estado de autenticación observado: {safe_text(best_auth.get('Resultado'))} | "
        f"Control: {safe_text(best_auth.get('Control'))}"
    )
    document.add_paragraph(
        f"Evidencia principal: {truncate(clean_evidence_for_report(best_auth.get('Evidencia')), 320)}",
        style="List Bullet",
    )

    deduped_pages = _dedupe_pages_for_report(pages or [])
    post_login_pages = [
        page for page in deduped_pages
        if safe_text(page.get("discovery_context")).lower() == "post_login"
        and _is_meaningful_page_for_report(page)
    ]
    new_post_login_pages = [page for page in post_login_pages if page.get("is_new_post_login")]
    display_post_login = new_post_login_pages or post_login_pages
    protected_keywords = ["admin", "dashboard", "backoffice", "private"]
    protected_pages = [
        page for page in display_post_login
        if safe_text(page.get("classification")).lower() in {
            "protected", "protected_redirect_to_auth", "admin_candidate", "api_candidate", "sensitive_candidate"
        }
        or any(
            token in safe_text(page.get("final_url") or page.get("url")).lower()
            for token in protected_keywords
        )
    ]

    document.add_paragraph(
        f"Superficie descubierta en contexto post-login: {len(display_post_login)} URL(s). "
        f"Rutas potencialmente sensibles/protegidas: {len(protected_pages)}."
    )

    if not display_post_login:
        document.add_paragraph(
            "No se observaron rutas nuevas relevantes tras autenticación; se omite detalle repetido para reducir ruido.",
            style="List Bullet",
        )

    for page in protected_pages[:BRIEF_MAX_AUTH_ROWS]:
        final_url = safe_text(page.get("final_url") or page.get("url"))
        classification = safe_text(page.get("classification", "sin clasificar"))
        status_code = safe_text(page.get("status_code", ""))
        document.add_paragraph(
            f"Ruta post-login: {truncate(final_url, 120)} | HTTP: {status_code} | Clasificación: {classification}",
            style="List Bullet",
        )

    post_login_results = [
        row for row in results or []
        if safe_text(row.get("Módulo")).strip() == "Discovery post-login"
    ]
    compact_post_login_results = []
    seen_post_login_events = set()
    for row in post_login_results:
        key = _result_dedupe_key(row)
        if key in seen_post_login_events:
            continue
        seen_post_login_events.add(key)
        compact_post_login_results.append(row)
    if compact_post_login_results and display_post_login:
        document.add_paragraph(
            f"Evidencias adicionales de discovery post-login registradas: {len(compact_post_login_results)} evento(s) netos."
        )
        for row in compact_post_login_results[:3]:
            document.add_paragraph(
                truncate(clean_evidence_for_report(row.get("Evidencia")), 260),
                style="List Bullet",
            )


def add_top_findings_table(document, findings):
    document.add_heading("5. Hallazgos prioritarios", 1)

    if not findings:
        document.add_paragraph("No se identificaron hallazgos prioritarios en el alcance automatizado.")
        return

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    headers = ["Sev.", "Categoría", "Control", "Conf.", "Evidencia resumida", "Recomendación"]
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        hdr_row.cells[i].text = h
    style_header_row(hdr_row)

    for item in _prioritized_findings(findings, BRIEF_MAX_FINDINGS):
        row = table.add_row().cells
        sev = item.get("Severidad", "")
        confidence = item.get("Confianza", "")
        row[0].text = sev
        row[1].text = item.get("Módulo", "")
        row[2].text = truncate(item.get("Control", ""), 90)
        row[3].text = f"{confidence:.2f}" if isinstance(confidence, (int, float)) else truncate(str(confidence), 10)
        row[4].text = truncate(clean_evidence_for_report(item.get("Evidencia", "")), 220)
        row[5].text = truncate(item.get("Recomendación", ""), 220)
        style_severity_cell(row[0], sev)


def add_execution_errors_table(document, errors):
    document.add_heading("6. Errores o limitaciones de ejecución", 1)

    if not errors:
        document.add_paragraph("No se registraron errores de ejecución en los módulos automatizados.")
        return

    document.add_paragraph(
        "Los siguientes controles no pudieron completarse correctamente. "
        "Estos errores no deben interpretarse como ausencia de vulnerabilidad."
    )

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["Módulo", "Control", "Descripción", "Evidencia"]

    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for item in errors[:BRIEF_MAX_ERROR_ROWS]:
        row = table.add_row().cells
        row[0].text = truncate(item.get("Módulo", ""), 80)
        row[1].text = truncate(item.get("Control", ""), 100)
        row[2].text = truncate(item.get("Descripción", ""), 220)
        row[3].text = truncate(item.get("Evidencia", ""), 260)


def add_cases_checked(document, results):
    document.add_heading("7. Casos comprobados", 1)

    document.add_paragraph(
        "La siguiente tabla resume los controles evaluados, diferenciando hallazgos, errores y comprobaciones sin evidencia."
    )

    grouped = defaultdict(list)

    for item in results:
        grouped[item.get("Módulo", "Sin módulo")].append(item)

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    headers = ["Categoría", "Comprobaciones", "Total", "Hallazgos", "Errores", "Resultado general"]

    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for module, items in grouped.items():
        findings = [x for x in items if is_finding(x)]
        errors = [x for x in items if is_error(x)]

        if findings:
            general = "Requiere revisión"
        elif errors:
            general = "Ejecución incompleta"
        else:
            general = "Sin hallazgos evidenciados"

        examples = [item.get("Control", "") for item in items[:4]]

        row = table.add_row().cells
        row[0].text = module
        row[1].text = truncate("; ".join(examples), 260)
        row[2].text = str(len(items))
        row[3].text = str(len(findings))
        row[4].text = str(len(errors))
        row[5].text = general


def add_test_path_summary(document, pages=None, pages_count=None):
    document.add_heading("8. Trazabilidad resumida de pruebas ejecutadas", 1)

    statements = []

    if pages_count is not None:
        statements.append(
            f"El crawler identificó {pages_count} URL(s) HTML dentro del dominio objetivo."
        )

    auth_pages = [
        page for page in pages or []
        if is_reportable_page(page)
        and page.get("classification") in ["auth", "registration", "protected_redirect_to_auth"]
    ]

    if auth_pages:
        for page in auth_pages[:10]:
            form_summary = get_form_detection_summary(page)
            forms_count = form_summary["forms_count"]
            statements.append(
                f"Durante discovery se identificó la ruta {page.get('url')} "
                f"clasificada como {page.get('classification')}, con código HTTP {page.get('status_code')} "
                f"y {forms_count} formulario(s) detectado(s) ({form_summary['form_type']})."
            )

            if forms_count == 0:
                statements.append(
                    "Al no detectarse formulario renderizado, se considera probable que el flujo de autenticación "
                    "esté renderizado en cliente o use endpoints API. Se recomienda análisis de JavaScript, "
                    "interceptación con navegador/headless y pruebas autenticadas."
                )

    statements.extend([
        "Se evaluaron cabeceras de seguridad HTTP como HSTS, CSP, X-Frame-Options, X-Content-Type-Options, Referrer-Policy y Permissions-Policy.",
        "Se analizaron cookies para identificar ausencia de atributos Secure, HttpOnly y SameSite.",
        "Se revisaron configuración CORS, métodos HTTP, exposición de APIs, JWT embebidos y estado TLS/HTTPS.",
        "Se comprobaron rutas comunes de recursos sensibles y directorios expuestos.",
        "Se analizaron formularios, parámetros y rutas descubiertas en el alcance del crawler.",
        "Las pruebas automatizadas se ejecutaron de forma no destructiva y los hallazgos deben validarse manualmente antes de considerarse explotabilidad confirmada.",
    ])

    for statement in statements:
        document.add_paragraph(statement, style="List Bullet")


def infer_payload_families(text):
    value = safe_text(text).lower()
    families = []

    rules = [
        ("XSS script/event", ["<script", "onerror", "onload", "svg", "javascript:"]),
        ("SQLi boolean/error", ["or 1=1", "boolean", "sql", "syntax", "error-based"]),
        ("SQLi time-based", ["sleep(", "benchmark(", "waitfor delay", "time-based"]),
        ("SQLi union", ["union select", "union-based"]),
        ("SSTI expression", ["{{", "${", "<%=", "{%", "template"]),
        ("Open redirect", ["redirect", "next=", "url=", "return=", "//"]),
        ("DOM source/sink", ["document.location", "innerhtml", "eval(", "sink", "source"]),
        ("SSRF internal/metadata", ["169.254.169.254", "localhost", "127.0.0.1", "file://", "gopher://"]),
        ("Path traversal", ["../", "..\\", "%2e%2e%2f", "/etc/passwd", "win.ini"]),
    ]

    for family, markers in rules:
        if any(marker in value for marker in markers):
            families.append(family)

    return families


def _phase_for_module(module_name):
    return MODULE_PHASE.get(safe_text(module_name), "Otros")


def add_attack_timeline_section(document, results):
    document.add_heading("9. Timeline táctico del atacante (correlación por fases)", 1)

    if not results:
        document.add_paragraph("No hay resultados para construir timeline de ejecución.")
        return

    grouped = defaultdict(list)
    for item in results:
        phase = _phase_for_module(item.get("Módulo", ""))
        grouped[phase].append(item)

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    headers = ["Fase", "Controles", "Hallazgos", "Errores", "Resumen"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    style_header_row(table.rows[0])

    for phase in PHASE_ORDER:
        items = grouped.get(phase, [])
        if not items:
            continue

        findings = [x for x in items if is_finding(x)]
        errors = [x for x in items if is_error(x)]

        if findings:
            summary = "Con evidencia de riesgo"
        elif errors:
            summary = "Cobertura parcial por errores"
        else:
            summary = "Sin explotación evidenciada"

        row = table.add_row().cells
        row[0].text = phase
        row[1].text = str(len(items))
        row[2].text = str(len(findings))
        row[3].text = str(len(errors))
        row[4].text = summary

    document.add_paragraph(
        "Hitos principales por fase (primeros controles relevantes o con hallazgo):"
    )
    for phase in PHASE_ORDER:
        items = grouped.get(phase, [])
        if not items:
            continue

        notable = [x for x in items if is_finding(x)] or items
        for item in notable[:3]:
            document.add_paragraph(
                f"{phase} | {item.get('Módulo', '')} | {truncate(item.get('Control', ''), 90)} | "
                f"Resultado: {item.get('Resultado', '')}",
                style="List Bullet",
            )


def summarize_module_status(items):
    statuses = {safe_text(x.get("Resultado")) for x in items}
    if any(status in FINDING_STATUSES for status in statuses):
        return "Con hallazgo"
    if any(status in {"Error", "No probado"} for status in statuses):
        return "Cobertura incompleta"
    return "Ejecutadas sin explotación"


def add_offensive_coverage_section(document, results):
    document.add_heading("10. Cobertura ofensiva — Explotación, protección y brechas", 1)

    document.add_paragraph(
        "Esta sección distingue TRES estados posibles por cada prueba ofensiva ejecutada:"
    )
    legend_items = [
        ("EXPLOTADO",   "Vulnerabilidad confirmada. El exploit funcionó."),
        ("POSIBLE",     "Indicador de vulnerabilidad encontrado. Requiere validación manual."),
        ("PROTEGIDO",   "Prueba ejecutada. El objetivo resistió el ataque. El control es efectivo."),
        ("NO CUBIERTO", "La prueba no se pudo completar (error de ejecución) o el módulo no corrió."),
    ]
    leg_table = document.add_table(rows=1, cols=2)
    leg_table.style = "Table Grid"
    leg_table.rows[0].cells[0].text = "Estado"
    leg_table.rows[0].cells[1].text = "Significado"
    style_header_row(leg_table.rows[0])
    for badge_key, meaning in legend_items:
        row = leg_table.add_row().cells
        row[0].text = badge_key
        row[1].text = meaning
        _apply_badge(row[0], badge_key)

    document.add_paragraph("")

    # Collect all offensive results
    offensive_results = [
        item for item in results
        if item.get("Módulo") in OFFENSIVE_MODULES
    ]

    # Detect modules that were NEVER run (no results at all)
    ran_modules = {item.get("Módulo") for item in offensive_results}
    never_run = [m for m in OFFENSIVE_MODULES if m not in ran_modules]

    if not offensive_results and never_run:
        document.add_paragraph(
            "⚠ NINGÚN módulo ofensivo se ejecutó en esta auditoría. "
            "La cobertura de explotación es CERO. Los siguientes módulos no se probaron:"
        )
        for m in never_run:
            document.add_paragraph(m, style="List Bullet")
        return

    # Per-module breakdown
    grouped = defaultdict(list)
    for item in offensive_results:
        grouped[item.get("Módulo", "Sin módulo")].append(item)

    for module in sorted(grouped.keys()):
        items = grouped[module]
        findings = [x for x in items if x.get("Resultado") in EXPLOIT_RESULT]
        protected = [x for x in items if x.get("Resultado") in PROTECTED_RESULT]
        errors = [x for x in items if x.get("Resultado") in UNCOVERED_RESULT]

        # Module header with overall state
        if findings:
            overall_badge = "EXPLOTADO" if any(
                x.get("Resultado") == "Hallazgo" for x in findings
            ) else "POSIBLE"
        elif errors and not protected:
            overall_badge = "NO CUBIERTO"
        elif protected:
            overall_badge = "PROTEGIDO"
        else:
            overall_badge = "NO CUBIERTO"

        document.add_heading(f"Módulo: {module}", 3)
        p = document.add_paragraph()
        bg_h, fg_h = COVERAGE_BADGE.get(overall_badge, ("616A6B", "FFFFFF"))
        run = p.add_run(f"  Estado global: {overall_badge}  ")
        run.font.color.rgb = RGBColor(
            int(fg_h[:2], 16), int(fg_h[2:4], 16), int(fg_h[4:], 16)
        )
        run.font.bold = True
        p.add_run(f"  — {len(items)} controles evaluados | "
                  f"{len(findings)} explotados | "
                  f"{len(protected)} protegidos | "
                  f"{len(errors)} sin cobertura")

        # Per-control table
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0]
        for i, h in enumerate(["Estado", "Control/Técnica", "Resultado", "Evidencia clave", "Recomendación"]):
            hdr.cells[i].text = h
        style_header_row(hdr)

        for item in items:
            res = safe_text(item.get("Resultado", ""))
            badge = coverage_state(res)
            row = table.add_row().cells
            row[0].text = badge
            row[1].text = truncate(item.get("Control", ""), 90)
            row[2].text = res
            # For exploited/possible: show full evidence; for protected: show concise note
            if badge in ("EXPLOTADO", "POSIBLE"):
                row[3].text = truncate(clean_evidence_for_report(item.get("Evidencia", "")), 280)
                row[4].text = truncate(item.get("Recomendación", ""), 200)
            else:
                row[3].text = truncate(clean_evidence_for_report(item.get("Evidencia", "")), 120) or "Sin evidencia de vulnerabilidad."
                row[4].text = "Control efectivo. Sin acción correctiva inmediata."
            _apply_badge(row[0], badge)

        document.add_paragraph("")

    # Modules never run = coverage gap warning
    if never_run:
        document.add_heading("⚠ Módulos ofensivos NO ejecutados (brechas de cobertura)", 3)
        document.add_paragraph(
            "Los siguientes módulos ofensivos NO se ejecutaron en esta auditoría. "
            "Su ausencia NO significa que el objetivo sea seguro frente a estos ataques — "
            "simplemente no se han probado. Deben incluirse en la próxima ejecución."
        )
        for m in never_run:
            document.add_paragraph(m, style="List Bullet")




def add_detailed_findings(document, findings):
    document.add_heading("11. Detalle técnico de hallazgos", 1)

    if not findings:
        document.add_paragraph("No se identificaron hallazgos técnicos relevantes.")
        return

    document.add_paragraph(
        "Cada hallazgo incluye: severidad, evidencia concreta obtenida, y recomendación accionable. "
        "Los hallazgos marcados como EXPLOTADO o POSIBLE requieren remediación antes de la siguiente auditoria."
    )

    for item in sort_results(findings):
        sev = item.get("Severidad", "Informativa")
        resultado = safe_text(item.get("Resultado", ""))
        badge = coverage_state(resultado)
        bg = SEV_BG.get(sev, "616A6B")

        heading = document.add_heading(
            f"[{sev}] {item.get('Módulo', '')} — {item.get('Control', '')}",
            2,
        )
        for run in heading.runs:
            r_int = int(bg[0:2], 16)
            g_int = int(bg[2:4], 16)
            b_int = int(bg[4:6], 16)
            run.font.color.rgb = RGBColor(r_int, g_int, b_int)
            run.font.bold = True

        # Compact info table for each finding
        tbl = document.add_table(rows=4, cols=2)
        tbl.style = "Table Grid"

        labels = ["Estado", "Descripción", "Evidencia", "Recomendación"]
        values = [
            resultado,
            truncate(item.get("Descripción", ""), 600),
            truncate(clean_evidence_for_report(item.get("Evidencia", "")), 800),
            truncate(item.get("Recomendación", ""), 600),
        ]
        for i, (label, value) in enumerate(zip(labels, values)):
            tbl.rows[i].cells[0].text = label
            tbl.rows[i].cells[1].text = value
            _set_cell_bg(tbl.rows[i].cells[0], HEADER_BG)
            _set_cell_font(tbl.rows[i].cells[0], HEADER_FG, bold=True)
            # Highlight estado cell
            if label == "Estado":
                _apply_badge(tbl.rows[i].cells[1], badge)

        document.add_paragraph("")


def add_discovery_dictionary_section(document, discovery):
    document.add_heading("Discovery activo mediante diccionario", 1)

    if not discovery:
        document.add_paragraph("No se proporcionaron resultados de discovery activo.")
        return

    discovered = discovery.get("discovered", [])
    metrics = discovery.get("metrics", {})

    normalized_discovered = []
    for item in discovered or []:
        if isinstance(item, dict):
            normalized_discovered.append(item)
            continue

        url = safe_text(item).strip()
        if not url:
            continue
        normalized_discovered.append({
            "source": "unknown",
            "requested_url": url,
            "final_url": url,
            "status_code": "",
            "classification": "",
            "observation": "",
        })

    document.add_paragraph(
        "Se ejecutó una fase de descubrimiento activo mediante diccionario de rutas comunes. "
        "El informe omite rutas inexistentes, soft-404 y ruido operativo, manteniendo únicamente rutas con valor técnico."
    )

    document.add_paragraph(f"URLs procesadas: {metrics.get('total_discovered', len(normalized_discovered))}")
    document.add_paragraph(f"Rutas relevantes reportables: {metrics.get('reportable_discovered', 0)}")
    document.add_paragraph(f"Rutas de autenticación/registro: {metrics.get('auth_routes', 0) + metrics.get('registration_routes', 0)}")
    document.add_paragraph(f"Rutas protegidas/API/sensibles: {metrics.get('protected_routes', 0) + metrics.get('api_candidates', 0) + metrics.get('sensitive_candidates', 0)}")
    document.add_paragraph(f"Soft-404 omitidos: {metrics.get('soft_404', 0)}")

    reportable_count = len([
        item for item in normalized_discovered
        if item.get("classification") in REPORTABLE_PAGE_CLASSES
        and str(item.get("status_code")) != "404"
    ])

    if reportable_count <= 0:
        document.add_paragraph(
            "No se identificaron rutas relevantes mediante diccionario. "
            "Las rutas inexistentes o soft-404 se han omitido del informe para reducir ruido."
        )
        return

    document.add_paragraph(
        "Detalle de URLs relevantes ya consolidado en la Sección 3 para evitar duplicidad entre tablas de discovery."
    )


def add_conclusion(document, findings, errors, results=None):
    document.add_heading("12. Conclusión y próximos pasos", 1)

    offensive_assurance = None
    for item in results or []:
        if item.get("Módulo") == "Aseguramiento ofensivo":
            offensive_assurance = item
            break

    if findings:
        document.add_paragraph(
            "La auditoría automatizada ha identificado hallazgos que requieren revisión técnica. "
            "Se recomienda priorizar los clasificados como críticos o altos y validar manualmente su explotabilidad."
        )
    elif errors:
        document.add_paragraph(
            "No se han identificado hallazgos confirmados, pero existen errores o limitaciones de ejecución en algunos módulos. "
            "La ausencia de hallazgos no debe interpretarse como ausencia de vulnerabilidades."
        )
    else:
        document.add_paragraph(
            "No se han identificado hallazgos relevantes en el alcance automatizado. "
            "Se recomienda complementar con pruebas manuales autenticadas y revisión de lógica de negocio."
        )

    if offensive_assurance:
        document.add_paragraph(
            "Estado de aseguramiento ofensivo: "
            f"{offensive_assurance.get('Resultado', '')}. "
            f"{offensive_assurance.get('Descripción', '')}"
        )

    document.add_paragraph("Plan recomendado:")
    document.add_paragraph("Validar manualmente cualquier hallazgo crítico o alto.", style="List Number")
    document.add_paragraph("Resolver errores de ejecución, bloqueos o limitaciones técnicas.", style="List Number")
    document.add_paragraph("Reejecutar la auditoría tras corregir limitaciones.", style="List Number")
    document.add_paragraph("Ampliar la auditoría con credenciales, navegador/headless y roles de usuario si aplica.", style="List Number")


def generate_word_report(
    audit_name: str,
    target_url: str,
    results: list,
    pages: list | None = None,
    discovery: dict | None = None,
    pages_count: int | None = None,
    scan_mode: str | None = None,
):
    os.makedirs("generated_reports", exist_ok=True)

    pages = _dedupe_pages_for_report(pages or [])
    if pages_count is None:
        pages_count = len(pages)

    document = Document()
    set_document_style(document)

    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    deduped_results = dedupe_results_for_report(results)
    findings = [x for x in deduped_results if is_finding(x)]
    errors = [x for x in deduped_results if is_error(x)]
    oks = [x for x in deduped_results if is_ok(x)]

    # Extract concrete intelligence before building sections
    assets = extract_sensitive_assets(deduped_results, pages)

    add_title_page(document, audit_name, target_url, scan_mode, pages_count)

    # ── Section 1: Executive summary ──────────────────────────────────
    document.add_heading("1. Resumen ejecutivo", 1)

    crit_high = [f for f in findings if f.get("Severidad") in ("Crítica", "Alta")]
    ports_found = assets.get("ports") or []
    users_found = assets.get("users") or []
    techs_found = assets.get("technologies") or []
    creds_found = assets.get("credentials") or []
    never_run = [m for m in OFFENSIVE_MODULES if m not in {
        item.get("Módulo") for item in deduped_results
    }]

    document.add_paragraph(
        f"Auditoría automatizada ejecutada sobre {target_url}. "
        f"Se han lanzado {len(deduped_results)} pruebas netas (sin duplicados) distribuidas en "
        f"{len(set(x.get('Módulo','') for x in deduped_results))} módulos."
    )

    # Highlight critical findings upfront
    if crit_high:
        document.add_paragraph(
            f"⚠ ATENCIÓN: Se identificaron {len(crit_high)} hallazgo(s) de severidad ALTA o CRÍTICA "
            f"que requieren remediación inmediata.",
        )
        for f in sort_results(crit_high)[:5]:
            document.add_paragraph(
                f"• [{f.get('Severidad','')}] {f.get('Módulo','')} — "
                f"{truncate(f.get('Control',''), 80)}: "
                f"{truncate(f.get('Evidencia',''), 120)}",
                style="List Bullet",
            )

    if ports_found:
        risky_ports = [p for p in ports_found if p.get("severity") == "Alta"]
        if risky_ports:
            document.add_paragraph(
                f"Puertos de alto riesgo expuestos: "
                f"{', '.join(str(p['port']) for p in risky_ports)}"
            )

    if techs_found:
        techs_str = ", ".join(f"{t['tech']}/{t['version']}" for t in techs_found[:6])
        document.add_paragraph(f"Tecnologías identificadas: {techs_str}")

    if users_found:
        document.add_paragraph(
            f"Usuarios/emails descubiertos: {len(users_found)} "
            f"({', '.join(users_found[:4])}{'...' if len(users_found) > 4 else ''})"
        )

    if creds_found:
        document.add_paragraph(
            f"⚠ CREDENCIALES O DATOS SENSIBLES encontrados: {len(creds_found)} indicador(es)."
        )

    if never_run:
        document.add_paragraph(
            f"⚠ COBERTURA INCOMPLETA: {len(never_run)} módulo(s) ofensivo(s) no ejecutados "
            f"({', '.join(never_run)}). Estos ataques NO se han probado."
        )

    add_summary_table(document, findings, errors, oks)
    add_severity_table(document, findings)
    add_c_level_one_page_summary(document, target_url, findings, errors, assets, deduped_results)

    # ── Section 2: Sensitive assets ───────────────────────────────────
    add_sensitive_assets_section(document, assets)

    # ── Section 3: Discovered surface ────────────────────────────────
    document.add_heading("3. Superficie relevante descubierta por crawler/discovery", 1)
    add_discovered_surface_body(document, pages)
    add_discovery_dictionary_section(document, discovery)

    # ── Section 4: Auth surface ───────────────────────────────────────
    add_auth_surface_summary(document, pages)
    add_authenticated_session_evidence(document, deduped_results, pages)

    # ── Section 5: Top findings ───────────────────────────────────────
    add_top_findings_table(document, findings)
    add_business_risk_matrix(document, findings)
    add_compliance_mapping(document, findings)
    add_manual_offensive_backlog(document, findings, pages, assets)
    add_defensive_playbook(document, assets, findings)
    add_asset_risk_scoring_section(document, findings, assets, deduped_results)

    # ── Section 6: Execution errors / limitations ────────────────────
    add_execution_errors_table(document, errors)

    # Keep report concise and decision-oriented: omit verbose appendix-style sections.

    # ── Section 7: Conclusion ─────────────────────────────────────────
    add_conclusion(document, findings, errors, deduped_results)

    safe_name = audit_name.replace(" ", "_").replace("/", "_").replace("\\", "_")
    file_path = f"generated_reports/{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    document.save(file_path)

    return file_path