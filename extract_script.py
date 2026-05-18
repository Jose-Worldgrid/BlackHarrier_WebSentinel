import docx
import sys
import re
import os

file_path = r'C:\Users\jcalderonserrano\Desktop\auditorias\Auditoria WEB\WEB\venv\webaudit-toolkit\generated_reports\Auditoría_Web_-_2026-05-18_20260518_124538.docx'
keywords = ['Nmap', 'Host', 'Puerto', 'glutenzero', '84.247.130.50']

try:
    doc = docx.Document(file_path)
    print(f'File: {file_path}')
    detected = set()
    for para in doc.paragraphs:
        if any(keyword.lower() in para.text.lower() for keyword in keywords):
            text = para.text.strip()
            if text and text not in detected:
                print(text)
                detected.add(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if any(keyword.lower() in cell.text.lower() for keyword in keywords):
                    text = cell.text.strip()
                    if text and text not in detected:
                        print(text)
                        detected.add(text)
except Exception as e:
    print(f'Error: {e}')
